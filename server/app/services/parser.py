import os
import pandas as pd
import pypdf
from docx import Document
from typing import Dict, Any, Tuple, List

class DataParser:
    @staticmethod
    def parse_file(file_path: str, file_type: str) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        """
        Parses CSV, XLSX, PDF, or DOCX files and returns a pandas DataFrame along with raw schema metadata.
        """
        file_type = file_type.lower()
        if file_type == 'csv':
            return DataParser._parse_csv(file_path)
        elif file_type == 'xlsx':
            return DataParser._parse_xlsx(file_path)
        elif file_type == 'pdf':
            return DataParser._parse_pdf(file_path)
        elif file_type == 'docx':
            return DataParser._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _parse_csv(file_path: str) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        # Read CSV with fallback encodings
        try:
            df = pd.read_csv(file_path)
        except UnicodeDecodeError:
            df = pd.read_csv(file_path, encoding='latin1')
        
        metadata = {
            "sheets": ["default"],
            "row_count": len(df),
            "column_count": len(df.columns)
        }
        return df, metadata

    @staticmethod
    def _parse_xlsx(file_path: str) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        xls = pd.ExcelFile(file_path)
        sheet_names = xls.sheet_names
        
        # Load the first sheet by default
        df = pd.read_excel(file_path, sheet_name=sheet_names[0])
        
        metadata = {
            "sheets": sheet_names,
            "row_count": len(df),
            "column_count": len(df.columns)
        }
        return df, metadata

    @staticmethod
    def _parse_pdf(file_path: str) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        """
        Extracts tabular data from PDF pages using pypdf.
        Falls back to page-by-page line extraction if no tables are matched.
        """
        reader = pypdf.PdfReader(file_path)
        text_lines = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                for line in text.split('\n'):
                    parts = [p.strip() for p in line.split(',') if p.strip()]
                    if not parts or len(parts) == 1:
                        parts = [p.strip() for p in line.split('\t') if p.strip()]
                    if not parts or len(parts) == 1:
                        parts = [p.strip() for p in line.split('  ') if p.strip()]
                    if parts and len(parts) > 1:
                        text_lines.append(parts)
                        
        if not text_lines:
            # Fallback to line extraction
            fallback_lines = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        if line.strip():
                            fallback_lines.append([f"Page {i+1}", line.strip()])
            if fallback_lines:
                df = pd.DataFrame(fallback_lines, columns=["Source_Page", "Extracted_Text"])
            else:
                df = pd.DataFrame([["Document", "No extractable text or tables found in this PDF."]], columns=["Source_Page", "Extracted_Text"])
        else:
            max_cols = max(len(row) for row in text_lines)
            padded_rows = [row + [None] * (max_cols - len(row)) for row in text_lines]
            df = pd.DataFrame(padded_rows)
            df.columns = [f"Col_{i}" for i in range(max_cols)]
            
        df = df.dropna(how='all').reset_index(drop=True)
        
        metadata = {
            "sheets": ["extracted_pdf_tables"],
            "row_count": len(df),
            "column_count": len(df.columns)
        }
        return df, metadata

    @staticmethod
    def _parse_docx(file_path: str) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        """
        Extracts table content from Microsoft Word files (.docx).
        """
        doc = Document(file_path)
        tables_data = []
        
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_cells)
            if table_rows:
                tables_data.append(table_rows)
                
        if not tables_data:
            # Fallback to paragraph parsing
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            df = pd.DataFrame(paragraphs, columns=["Paragraphs"])
        else:
            # Use the first table
            first_table = tables_data[0]
            if len(first_table) > 1:
                headers = first_table[0]
                rows = first_table[1:]
                headers = [h if h else f"Col_{i}" for i, h in enumerate(headers)]
                df = pd.DataFrame(rows, columns=headers)
            else:
                df = pd.DataFrame(first_table)
                
        metadata = {
            "sheets": ["extracted_docx_tables"],
            "row_count": len(df),
            "column_count": len(df.columns)
        }
        return df, metadata
